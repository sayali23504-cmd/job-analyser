from docx import Document


def build_resume_docx(path: str, job_title: str, company: str, summary_points: list[str], ats_keywords: list[str]) -> None:
    doc = Document()
    doc.add_heading("Tailored Resume Draft", level=1)
    doc.add_paragraph(f"Target Role: {job_title}")
    doc.add_paragraph(f"Target Company: {company}")

    doc.add_heading("Tailoring Suggestions", level=2)
    for point in summary_points:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("ATS Keyword Focus", level=2)
    for kw in ats_keywords[:30]:
        doc.add_paragraph(kw, style="List Bullet")

    doc.add_paragraph("Note: This draft does not invent experience or credentials.")
    doc.save(path)
